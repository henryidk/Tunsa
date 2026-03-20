from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# --- Estilos globales ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

# Colores
COLOR_TITULO_DOC = RGBColor(0x1E, 0x3A, 0x5F)   # azul oscuro
COLOR_TITULO_CU  = RGBColor(0x1F, 0x4E, 0x79)   # azul medio
COLOR_LABEL      = RGBColor(0x2E, 0x75, 0xB6)   # azul etiqueta
COLOR_SECCION    = RGBColor(0x2E, 0x75, 0xB6)   # azul sección
COLOR_SEPARADOR  = RGBColor(0xBD, 0xD7, 0xEE)   # azul claro línea

def set_font(run, bold=False, italic=False, size=None, color=None, name='Calibri'):
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    run.font.name = name

def add_titulo_documento(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    set_font(run, bold=True, size=18, color=COLOR_TITULO_DOC)
    p.paragraph_format.space_after = Pt(4)

def add_subtitulo_documento(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    set_font(run, size=11, color=RGBColor(0x60, 0x60, 0x60))
    p.paragraph_format.space_after = Pt(20)

def add_titulo_cu(doc, texto):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    set_font(run, bold=True, size=13, color=COLOR_TITULO_CU)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)

def add_campo(doc, label, valor):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    r1 = p.add_run(label + ": ")
    set_font(r1, bold=True, color=COLOR_LABEL, size=11)
    r2 = p.add_run(valor)
    set_font(r2, size=11)

def add_seccion(doc, titulo):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(titulo)
    set_font(run, bold=True, size=11, color=COLOR_SECCION)

def add_item_lista(doc, texto, nivel=0):
    p = doc.add_paragraph(style='List Bullet' if nivel == 0 else 'List Bullet 2')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.3 + nivel * 0.2)
    run = p.add_run(texto)
    set_font(run, size=11)

def add_item_numerado(doc, texto):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(texto)
    set_font(run, size=11)

def add_separador(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'BDD7EE')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_titulo_seccion_mayor(doc, texto):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    set_font(run, bold=True, size=14, color=COLOR_TITULO_DOC)
    p.paragraph_format.space_after = Pt(10)

# =========================================================
# PORTADA
# =========================================================
add_titulo_documento(doc, "Casos de Uso del Sistema")
add_titulo_documento(doc, "Proyecto Tunsa")
add_subtitulo_documento(doc, "Sistema de Gestión de Rentas de Equipos\nMarzo 2026")

# =========================================================
# SECCIÓN 1 — AUTENTICACIÓN Y SEGURIDAD
# =========================================================
add_titulo_seccion_mayor(doc, "Módulo 1 — Autenticación y Seguridad")

# CU-01
add_titulo_cu(doc, "CU-01: Iniciar Sesión")
add_campo(doc, "Nombre", "Iniciar Sesión")
add_campo(doc, "Código", "CU-01")
add_campo(doc, "Actor principal", "Todos los usuarios (Administrador, Secretaria, Colaborador, Encargado de Máquinas)")
add_campo(doc, "Descripción", "Permite a un usuario autenticarse en el sistema mediante credenciales para acceder a las funciones según su rol.")

add_seccion(doc, "Precondiciones")
add_item_lista(doc, "El usuario debe tener una cuenta activa en el sistema.")
add_item_lista(doc, "El sistema debe estar disponible.")

add_seccion(doc, "Flujo Básico")
add_item_numerado(doc, "El usuario accede a la pantalla de login.")
add_item_numerado(doc, "Ingresa su nombre de usuario y contraseña.")
add_item_numerado(doc, "El sistema valida las credenciales contra la base de datos.")
add_item_numerado(doc, "El sistema genera un accessToken (15 min) y un refreshToken (7 días).")
add_item_numerado(doc, "El sistema registra el evento LOGIN_SUCCESS en el AuditLog con IP y user agent.")
add_item_numerado(doc, "Si mustChangePassword = true, redirige al formulario de cambio de contraseña.")
add_item_numerado(doc, "El sistema redirige al dashboard correspondiente según el rol.")

add_seccion(doc, "Flujo Alterno")
add_item_lista(doc, "FA1 – Credenciales inválidas: El sistema muestra un mensaje de error y registra LOGIN_FAILED en el AuditLog.")
add_item_lista(doc, "FA2 – Usuario inactivo: El sistema rechaza el acceso con mensaje de cuenta desactivada.")
add_item_lista(doc, "FA3 – Demasiados intentos: El rate limiter bloquea temporalmente la IP.")

add_seccion(doc, "Postcondiciones")
add_item_lista(doc, "El usuario queda autenticado con tokens activos en sesión.")
add_item_lista(doc, "Se registra el acceso en el AuditLog.")

add_separador(doc)

# CU-02
add_titulo_cu(doc, "CU-02: Cerrar Sesión")
add_campo(doc, "Nombre", "Cerrar Sesión")
add_campo(doc, "Código", "CU-02")
add_campo(doc, "Actor principal", "Todos los usuarios autenticados")
add_campo(doc, "Descripción", "Permite al usuario finalizar su sesión activa, revocando todos sus tokens de seguridad.")

add_seccion(doc, "Precondiciones")
add_item_lista(doc, "El usuario debe estar autenticado.")

add_seccion(doc, "Flujo Básico")
add_item_numerado(doc, "El usuario hace clic en el botón de cerrar sesión.")
add_item_numerado(doc, "El sistema revoca todos los refreshTokens activos del usuario en la base de datos.")
add_item_numerado(doc, "El sistema elimina la cookie httpOnly del refreshToken.")
add_item_numerado(doc, "El sistema elimina el estado de autenticación del store (Zustand).")
add_item_numerado(doc, "El sistema registra LOGOUT en el AuditLog.")
add_item_numerado(doc, "El sistema redirige al usuario a la pantalla de login.")

add_seccion(doc, "Flujo Alterno")
add_item_lista(doc, "FA1 – Error de red: El sistema cierra la sesión localmente aunque falle la revocación en el servidor.")

add_seccion(doc, "Postcondiciones")
add_item_lista(doc, "La sesión queda invalidada en el servidor.")
add_item_lista(doc, "El usuario es redirigido al login.")

add_separador(doc)

# CU-03
add_titulo_cu(doc, "CU-03: Cambiar Contraseña (Forzada)")
add_campo(doc, "Nombre", "Cambiar Contraseña (Forzada)")
add_campo(doc, "Código", "CU-03")
add_campo(doc, "Actor principal", "Todos los usuarios")
add_campo(doc, "Descripción", "Obliga al usuario a establecer una nueva contraseña en su primer acceso o después de un reseteo administrativo, antes de poder usar el sistema.")

add_seccion(doc, "Precondiciones")
add_item_lista(doc, "El usuario debe haberse autenticado exitosamente.")
add_item_lista(doc, "El campo mustChangePassword debe ser true.")

add_seccion(doc, "Flujo Básico")
add_item_numerado(doc, "El sistema detecta mustChangePassword = true tras el login.")
add_item_numerado(doc, "El sistema muestra el modal CambiarPasswordModal.")
add_item_numerado(doc, "El usuario ingresa su nueva contraseña y la confirma.")
add_item_numerado(doc, "El sistema valida que la nueva contraseña no sea igual a la actual.")
add_item_numerado(doc, "El sistema actualiza la contraseña con hash bcrypt.")
add_item_numerado(doc, "El sistema establece mustChangePassword = false.")
add_item_numerado(doc, "El sistema redirige al dashboard correspondiente al rol.")

add_seccion(doc, "Flujo Alterno")
add_item_lista(doc, "FA1 – Contraseña igual a la actual: El sistema rechaza el cambio con mensaje de error.")
add_item_lista(doc, "FA2 – Contraseñas no coinciden: El sistema muestra error de validación.")

add_seccion(doc, "Postcondiciones")
add_item_lista(doc, "La contraseña queda actualizada en la base de datos.")
add_item_lista(doc, "mustChangePassword = false para el usuario.")
add_item_lista(doc, "El usuario accede al sistema.")

# =========================================================
# SECCIÓN 2 — GESTIÓN DE USUARIOS
# =========================================================
add_titulo_seccion_mayor(doc, "Módulo 2 — Gestión de Usuarios")

# CU-04
add_titulo_cu(doc, "CU-04: Crear Usuario")
add_campo(doc, "Nombre", "Crear Usuario")
add_campo(doc, "Código", "CU-04")
add_campo(doc, "Actor principal", "Administrador")
add_campo(doc, "Descripción", "El administrador registra un nuevo usuario en el sistema asignándole un rol y generando una contraseña temporal segura.")

add_seccion(doc, "Precondiciones")
add_item_lista(doc, "El actor debe estar autenticado con rol admin.")
add_item_lista(doc, "El username a registrar no debe existir en el sistema.")

add_seccion(doc, "Flujo Básico")
add_item_numerado(doc, 'El administrador accede a la sección Usuarios y hace clic en "Agregar usuario".')
add_item_numerado(doc, "El sistema muestra el modal AgregarUsuarioModal.")
add_item_numerado(doc, "El administrador completa: nombre, username, teléfono (opcional) y rol.")
add_item_numerado(doc, "El sistema genera una contraseña segura de 8 caracteres (mayúsculas, minúsculas, dígitos, especiales).")
add_item_numerado(doc, "El sistema crea el usuario con mustChangePassword = true e isActive = true.")
add_item_numerado(doc, "El modal muestra la contraseña temporal generada para que el administrador la copie.")
add_item_numerado(doc, "El sistema muestra un toast de éxito y actualiza la tabla de usuarios.")

add_seccion(doc, "Flujo Alterno")
add_item_lista(doc, 'FA1 – Username duplicado: El sistema muestra error "El nombre de usuario ya existe".')
add_item_lista(doc, "FA2 – Campos inválidos: El sistema muestra mensajes de validación por campo.")

add_seccion(doc, "Postcondiciones")
add_item_lista(doc, "El usuario queda registrado en la base de datos con rol asignado.")
add_item_lista(doc, "La contraseña temporal debe ser entregada al usuario por el administrador.")
add_item_lista(doc, "El usuario deberá cambiar su contraseña en el primer acceso.")

add_separador(doc)

# CU-05
add_titulo_cu(doc, "CU-05: Editar Usuario")
add_campo(doc, "Nombre", "Editar Usuario")
add_campo(doc, "Código", "CU-05")
add_campo(doc, "Actor principal", "Administrador")
add_campo(doc, "Descripción", "El administrador modifica los datos personales (nombre, username, teléfono) de un usuario existente.")

add_seccion(doc, "Precondiciones")
add_item_lista(doc, "El actor debe estar autenticado con rol admin.")
add_item_lista(doc, "El usuario a editar debe existir y ser diferente al usuario actual (no puede editarse a sí mismo).")

add_seccion(doc, "Flujo Básico")
add_item_numerado(doc, 'El administrador localiza el usuario en la tabla y hace clic en "Editar".')
add_item_numerado(doc, "El sistema muestra el modal EditarUsuarioModal con los datos actuales.")
add_item_numerado(doc, "El administrador modifica los campos deseados.")
add_item_numerado(doc, "El sistema valida los datos y actualiza el registro.")
add_item_numerado(doc, "El sistema muestra un toast de éxito y refresca la tabla.")

add_seccion(doc, "Flujo Alterno")
add_item_lista(doc, "FA1 – Username ya en uso: El sistema muestra error de unicidad.")
add_item_lista(doc, "FA2 – Sin cambios: El sistema acepta la solicitud sin modificar nada.")

add_seccion(doc, "Postcondiciones")
add_item_lista(doc, "Los datos del usuario quedan actualizados en la base de datos.")

add_separador(doc)

# CU-06
add_titulo_cu(doc, "CU-06: Desactivar Usuario")
add_campo(doc, "Nombre", "Desactivar Usuario")
add_campo(doc, "Código", "CU-06")
add_campo(doc, "Actor principal", "Administrador")
add_campo(doc, "Descripción", "El administrador desactiva la cuenta de un usuario, impidiendo su acceso al sistema e invalidando todas sus sesiones activas.")

add_seccion(doc, "Precondiciones")
add_item_lista(doc, "El actor debe estar autenticado con rol admin.")
add_item_lista(doc, "El usuario a desactivar debe estar activo y ser diferente al usuario actual.")

add_seccion(doc, "Flujo Básico")
add_item_numerado(doc, 'El administrador hace clic en "Desactivar" junto al usuario.')
add_item_numerado(doc, "El sistema muestra el modal de confirmación ConfirmDesactivarModal.")
add_item_numerado(doc, "El administrador confirma la acción.")
add_item_numerado(doc, "El sistema establece isActive = false en el usuario.")
add_item_numerado(doc, "El sistema revoca todos los refreshTokens activos del usuario.")
add_item_numerado(doc, "El sistema muestra un toast de éxito.")
add_item_numerado(doc, 'El estado del usuario en la tabla cambia a "Inactivo".')

add_seccion(doc, "Flujo Alterno")
add_item_lista(doc, "FA1 – El administrador cancela: No se realiza ningún cambio.")

add_seccion(doc, "Postcondiciones")
add_item_lista(doc, "El usuario no puede iniciar sesión.")
add_item_lista(doc, "Sus sesiones activas son invalidadas inmediatamente.")

add_separador(doc)

# CU-07
add_titulo_cu(doc, "CU-07: Activar Usuario")
add_campo(doc, "Nombre", "Activar Usuario")
add_campo(doc, "Código", "CU-07")
add_campo(doc, "Actor principal", "Administrador")
add_campo(doc, "Descripción", "El administrador reactiva la cuenta de un usuario previamente desactivado, permitiéndole volver a iniciar sesión.")

add_seccion(doc, "Precondiciones")
add_item_lista(doc, "El actor debe estar autenticado con rol admin.")
add_item_lista(doc, "El usuario a activar debe estar en estado inactivo.")

add_seccion(doc, "Flujo Básico")
add_item_numerado(doc, "El administrador localiza al usuario inactivo en la tabla.")
add_item_numerado(doc, 'Hace clic en el botón "Activar".')
add_item_numerado(doc, "El sistema establece isActive = true.")
add_item_numerado(doc, "El sistema muestra un toast de éxito y actualiza la tabla.")

add_seccion(doc, "Flujo Alterno")
add_item_lista(doc, "FA1 – Error del servidor: El sistema muestra toast de error y revierte el estado visual.")

add_seccion(doc, "Postcondiciones")
add_item_lista(doc, "El usuario puede volver a iniciar sesión con sus credenciales actuales.")

add_separador(doc)

# CU-08
add_titulo_cu(doc, "CU-08: Resetear Contraseña de Usuario")
add_campo(doc, "Nombre", "Resetear Contraseña de Usuario")
add_campo(doc, "Código", "CU-08")
add_campo(doc, "Actor principal", "Administrador")
add_campo(doc, "Descripción", "El administrador restablece la contraseña de un usuario, generando una nueva temporal y revocando todas sus sesiones activas.")

add_seccion(doc, "Precondiciones")
add_item_lista(doc, "El actor debe estar autenticado con rol admin.")
add_item_lista(doc, "El usuario objetivo debe existir en el sistema.")

add_seccion(doc, "Flujo Básico")
add_item_numerado(doc, 'El administrador hace clic en "Resetear" junto al usuario.')
add_item_numerado(doc, "El sistema muestra el modal ResetPasswordModal.")
add_item_numerado(doc, "El sistema genera una nueva contraseña segura de 8 caracteres.")
add_item_numerado(doc, "El sistema actualiza la contraseña con hash bcrypt.")
add_item_numerado(doc, "El sistema establece mustChangePassword = true.")
add_item_numerado(doc, "El sistema revoca todos los refreshTokens activos del usuario.")
add_item_numerado(doc, "El sistema registra PASSWORD_RESET_BY_ADMIN en el AuditLog.")
add_item_numerado(doc, "El modal muestra la nueva contraseña temporal.")

add_seccion(doc, "Flujo Alterno")
add_item_lista(doc, "FA1 – Error de servidor: El sistema muestra toast de error.")

add_seccion(doc, "Postcondiciones")
add_item_lista(doc, "La contraseña del usuario queda actualizada.")
add_item_lista(doc, "El usuario deberá cambiar su contraseña en su próximo acceso.")
add_item_lista(doc, "Todas las sesiones previas del usuario son invalidadas.")

# =========================================================
# SECCIÓN 3 — GESTIÓN DE EQUIPOS
# =========================================================
add_titulo_seccion_mayor(doc, "Módulo 3 — Gestión de Equipos")

# CU-09
add_titulo_cu(doc, "CU-09: Agregar Equipo al Inventario")
add_campo(doc, "Nombre", "Agregar Equipo al Inventario")
add_campo(doc, "Código", "CU-09")
add_campo(doc, "Actor principal", "Administrador")
add_campo(doc, "Descripción", "El administrador registra un nuevo equipo en el inventario con sus datos técnicos y precios de renta.")

add_seccion(doc, "Precondiciones")
add_item_lista(doc, "El actor debe estar autenticado con rol admin.")
add_item_lista(doc, "La numeración del equipo no debe existir previamente.")

add_seccion(doc, "Flujo Básico")
add_item_numerado(doc, "El administrador accede a la sección Equipos y abre AgregarEquipoModal.")
add_item_numerado(doc, "Completa: numeración, descripción, categoría, serie (opcional), cantidad, fecha de compra, monto de compra, tipo (Liviana / Pesada / Uso Propio).")
add_item_numerado(doc, "Opcionalmente ingresa precios de renta por día, semana y mes.")
add_item_numerado(doc, "El sistema valida los datos y crea el registro con isActive = true.")
add_item_numerado(doc, "El sistema muestra toast de éxito y actualiza la tabla.")

add_seccion(doc, "Flujo Alterno")
add_item_lista(doc, "FA1 – Numeración duplicada: El sistema rechaza con mensaje de error.")
add_item_lista(doc, "FA2 – Campos obligatorios vacíos: El sistema muestra validaciones por campo.")

add_seccion(doc, "Postcondiciones")
add_item_lista(doc, "El equipo queda registrado en el inventario como activo.")
add_item_lista(doc, "Los valores del inventario (total Q) se actualizan en las estadísticas.")

add_separador(doc)

# CU-10
add_titulo_cu(doc, "CU-10: Editar Datos de Equipo")
add_campo(doc, "Nombre", "Editar Datos de Equipo")
add_campo(doc, "Código", "CU-10")
add_campo(doc, "Actor principal", "Administrador")
add_campo(doc, "Descripción", "El administrador modifica los datos técnicos de un equipo existente, quedando registrado el historial de cambios en la bitácora.")

add_seccion(doc, "Precondiciones")
add_item_lista(doc, "El actor debe estar autenticado con rol admin.")
add_item_lista(doc, "El equipo debe estar activo en el inventario.")

add_seccion(doc, "Flujo Básico")
add_item_numerado(doc, 'El administrador hace clic en "Editar" en la fila del equipo.')
add_item_numerado(doc, "El sistema muestra EditarEquipoModal con los datos actuales.")
add_item_numerado(doc, "El administrador modifica los campos deseados.")
add_item_numerado(doc, "El sistema calcula las diferencias campo por campo.")
add_item_numerado(doc, "Para cada campo modificado, el sistema genera una entrada en la Bitácora con valor anterior y nuevo.")
add_item_numerado(doc, "El sistema actualiza el registro del equipo.")
add_item_numerado(doc, "El sistema muestra toast de éxito.")

add_seccion(doc, "Flujo Alterno")
add_item_lista(doc, "FA1 – Sin cambios: El sistema no genera entradas en bitácora.")
add_item_lista(doc, "FA2 – Numeración duplicada: El sistema rechaza el cambio.")

add_seccion(doc, "Postcondiciones")
add_item_lista(doc, "Los datos del equipo quedan actualizados.")
add_item_lista(doc, "Cada campo modificado queda registrado en la Bitácora con quién realizó el cambio.")

add_separador(doc)

# CU-11
add_titulo_cu(doc, "CU-11: Actualizar Precios de Renta de Equipo")
add_campo(doc, "Nombre", "Actualizar Precios de Renta de Equipo")
add_campo(doc, "Código", "CU-11")
add_campo(doc, "Actor principal", "Administrador")
add_campo(doc, "Descripción", "El administrador actualiza los precios de renta (día, semana, mes) de un equipo, con registro en bitácora.")

add_seccion(doc, "Precondiciones")
add_item_lista(doc, "El actor debe estar autenticado con rol admin.")
add_item_lista(doc, "El equipo debe existir y estar activo.")

add_seccion(doc, "Flujo Básico")
add_item_numerado(doc, 'El administrador hace clic en "Precios" en la fila del equipo.')
add_item_numerado(doc, "El sistema muestra PreciosEquipoModal con los precios actuales.")
add_item_numerado(doc, "El administrador actualiza uno o más precios.")
add_item_numerado(doc, "El sistema genera entradas en Bitácora para cada precio modificado.")
add_item_numerado(doc, "El sistema actualiza los precios y muestra toast de éxito.")

add_seccion(doc, "Flujo Alterno")
add_item_lista(doc, "FA1 – Precios negativos: El sistema muestra error de validación.")

add_seccion(doc, "Postcondiciones")
add_item_lista(doc, "Los nuevos precios se aplican a futuras solicitudes de renta.")
add_item_lista(doc, "Los cambios quedan registrados en Bitácora.")

add_separador(doc)

# CU-12
add_titulo_cu(doc, "CU-12: Dar de Baja Equipo")
add_campo(doc, "Nombre", "Dar de Baja Equipo")
add_campo(doc, "Código", "CU-12")
add_campo(doc, "Actor principal", "Administrador")
add_campo(doc, "Descripción", "El administrador retira un equipo del inventario activo indicando el motivo, sin eliminarlo del sistema.")

add_seccion(doc, "Precondiciones")
add_item_lista(doc, "El actor debe estar autenticado con rol admin.")
add_item_lista(doc, "El equipo debe estar activo.")

add_seccion(doc, "Flujo Básico")
add_item_numerado(doc, 'El administrador hace clic en "Dar de baja" en la fila del equipo.')
add_item_numerado(doc, "El sistema muestra BajaEquipoModal.")
add_item_numerado(doc, "El administrador ingresa el motivo de la baja.")
add_item_numerado(doc, "El sistema establece isActive = false, registra motivoBaja y fechaBaja.")
add_item_numerado(doc, "El sistema genera entrada en Bitácora.")
add_item_numerado(doc, 'El equipo se mueve a la pestaña "Dados de baja".')

add_seccion(doc, "Flujo Alterno")
add_item_lista(doc, "FA1 – Motivo vacío: El sistema requiere ingresar motivo antes de confirmar.")

add_seccion(doc, "Postcondiciones")
add_item_lista(doc, "El equipo queda en estado inactivo con fecha y motivo de baja.")
add_item_lista(doc, 'Aparece en la pestaña "Dados de baja" y no en el inventario activo.')

add_separador(doc)

# CU-13
add_titulo_cu(doc, "CU-13: Reactivar Equipo")
add_campo(doc, "Nombre", "Reactivar Equipo")
add_campo(doc, "Código", "CU-13")
add_campo(doc, "Actor principal", "Administrador")
add_campo(doc, "Descripción", "El administrador reactiva un equipo previamente dado de baja, reincorporándolo al inventario activo.")

add_seccion(doc, "Precondiciones")
add_item_lista(doc, "El actor debe estar autenticado con rol admin.")
add_item_lista(doc, "El equipo debe estar en estado de baja (isActive = false).")

add_seccion(doc, "Flujo Básico")
add_item_numerado(doc, 'El administrador accede a la pestaña "Dados de baja" en Equipos.')
add_item_numerado(doc, 'Hace clic en "Reactivar" en la fila del equipo (muestra spinner durante la operación).')
add_item_numerado(doc, "El sistema establece isActive = true y limpia motivoBaja y fechaBaja.")
add_item_numerado(doc, "El sistema genera entrada en Bitácora.")
add_item_numerado(doc, 'El equipo se mueve a la pestaña "Activos".')

add_seccion(doc, "Flujo Alterno")
add_item_lista(doc, "FA1 – Error de servidor: El sistema revierte el estado visual y muestra error.")

add_seccion(doc, "Postcondiciones")
add_item_lista(doc, "El equipo queda disponible en el inventario activo.")

add_separador(doc)

# CU-14
add_titulo_cu(doc, "CU-14: Generar Reporte PDF de Inventario")
add_campo(doc, "Nombre", "Generar Reporte PDF de Inventario de Equipos")
add_campo(doc, "Código", "CU-14")
add_campo(doc, "Actor principal", "Administrador")
add_campo(doc, "Descripción", "El administrador genera y descarga un reporte PDF con el inventario completo de equipos, organizado por tipo con valores y estadísticas.")

add_seccion(doc, "Precondiciones")
add_item_lista(doc, "El actor debe estar autenticado con rol admin.")
add_item_lista(doc, "Deben existir equipos registrados en el sistema.")

add_seccion(doc, "Flujo Básico")
add_item_numerado(doc, 'El administrador hace clic en "Generar reporte" en la sección Equipos.')
add_item_numerado(doc, "El sistema recopila todos los equipos de la base de datos.")
add_item_numerado(doc, "El sistema genera el PDF en formato A4 landscape con: tarjetas de resumen (conteo y valor por tipo), tablas por sección (Liviana, Pesada, Uso Propio, Dados de Baja), encabezados y pies de página con número de página, y formato de moneda (Q).")
add_item_numerado(doc, "El sistema descarga el archivo como inventario_equipos_YYYY-MM-DD.pdf.")

add_seccion(doc, "Flujo Alterno")
add_item_lista(doc, "FA1 – Sin equipos: El sistema genera un reporte vacío o con mensaje informativo.")

add_seccion(doc, "Postcondiciones")
add_item_lista(doc, "El archivo PDF queda descargado en el dispositivo del usuario.")

# =========================================================
# SECCIÓN 4 — RENTAS Y CLIENTES
# =========================================================
add_titulo_seccion_mayor(doc, "Módulo 4 — Rentas y Clientes")

# CU-15
add_titulo_cu(doc, "CU-15: Registrar Cliente")
add_campo(doc, "Nombre", "Registrar Cliente")
add_campo(doc, "Código", "CU-15")
add_campo(doc, "Actor principal", "Administrador, Secretaria")
add_campo(doc, "Descripción", "El actor registra un nuevo cliente en el sistema con sus datos de identificación para gestionar sus rentas.")

add_seccion(doc, "Precondiciones")
add_item_lista(doc, "El actor debe estar autenticado con rol admin o secretaria.")
add_item_lista(doc, "El DPI del cliente no debe estar registrado previamente.")

add_seccion(doc, "Flujo Básico")
add_item_numerado(doc, 'El actor accede a la sección Clientes y hace clic en "Registrar cliente".')
add_item_numerado(doc, "El sistema muestra RegistrarClienteModal.")
add_item_numerado(doc, "El actor ingresa: nombre, DPI, teléfono (opcional).")
add_item_numerado(doc, "El sistema genera automáticamente un código único con formato CLI-XXXX.")
add_item_numerado(doc, "El sistema crea el registro del cliente.")
add_item_numerado(doc, "El sistema muestra toast de éxito y actualiza la tabla.")

add_seccion(doc, "Flujo Alterno")
add_item_lista(doc, "FA1 – DPI duplicado: El sistema muestra error de DPI ya registrado.")
add_item_lista(doc, "FA2 – Nombre o DPI vacíos: El sistema muestra validaciones.")

add_seccion(doc, "Postcondiciones")
add_item_lista(doc, "El cliente queda registrado con código único asignado.")
add_item_lista(doc, "Puede ser asociado a solicitudes de renta.")

add_separador(doc)

# CU-16
add_titulo_cu(doc, "CU-16: Editar Cliente")
add_campo(doc, "Nombre", "Editar Cliente")
add_campo(doc, "Código", "CU-16")
add_campo(doc, "Actor principal", "Administrador, Secretaria")
add_campo(doc, "Descripción", "El actor actualiza los datos de un cliente registrado en el sistema.")

add_seccion(doc, "Precondiciones")
add_item_lista(doc, "El actor debe estar autenticado con rol admin o secretaria.")
add_item_lista(doc, "El cliente debe existir en el sistema.")

add_seccion(doc, "Flujo Básico")
add_item_numerado(doc, "El actor localiza al cliente y hace clic en editar.")
add_item_numerado(doc, "El sistema muestra el formulario con los datos actuales.")
add_item_numerado(doc, "El actor modifica los datos deseados (nombre, teléfono).")
add_item_numerado(doc, "El sistema valida y actualiza el registro.")
add_item_numerado(doc, "El sistema muestra toast de éxito.")

add_seccion(doc, "Flujo Alterno")
add_item_lista(doc, "FA1 – DPI modificado a uno existente: El sistema rechaza con error de duplicado.")

add_seccion(doc, "Postcondiciones")
add_item_lista(doc, "Los datos del cliente quedan actualizados en la base de datos.")

add_separador(doc)

# CU-17
add_titulo_cu(doc, "CU-17: Eliminar Cliente")
add_campo(doc, "Nombre", "Eliminar Cliente")
add_campo(doc, "Código", "CU-17")
add_campo(doc, "Actor principal", "Administrador")
add_campo(doc, "Descripción", "El administrador elimina permanentemente el registro de un cliente del sistema.")

add_seccion(doc, "Precondiciones")
add_item_lista(doc, "El actor debe estar autenticado con rol admin.")
add_item_lista(doc, "El cliente debe existir en el sistema.")

add_seccion(doc, "Flujo Básico")
add_item_numerado(doc, "El administrador selecciona la opción de eliminar para un cliente.")
add_item_numerado(doc, "El sistema solicita confirmación.")
add_item_numerado(doc, "El administrador confirma la eliminación.")
add_item_numerado(doc, "El sistema elimina el registro del cliente.")
add_item_numerado(doc, "El sistema actualiza la tabla de clientes.")

add_seccion(doc, "Flujo Alterno")
add_item_lista(doc, "FA1 – El administrador cancela: No se realiza ningún cambio.")
add_item_lista(doc, "FA2 – Cliente con rentas activas: El sistema rechaza la eliminación con mensaje de advertencia.")

add_seccion(doc, "Postcondiciones")
add_item_lista(doc, "El registro del cliente es eliminado permanentemente.")

add_separador(doc)

# CU-18
add_titulo_cu(doc, "CU-18: Crear Solicitud de Renta")
add_campo(doc, "Nombre", "Crear Solicitud de Renta")
add_campo(doc, "Código", "CU-18")
add_campo(doc, "Actor principal", "Encargado de Máquinas")
add_campo(doc, "Descripción", "El encargado de máquinas crea una solicitud de renta de equipos para un cliente, la cual queda pendiente de aprobación.")

add_seccion(doc, "Precondiciones")
add_item_lista(doc, "El actor debe estar autenticado con rol encargado_maquinas.")
add_item_lista(doc, "Deben existir clientes y equipos disponibles registrados en el sistema.")

add_seccion(doc, "Flujo Básico")
add_item_numerado(doc, "El actor accede a Solicitudes de Renta y crea una nueva solicitud.")
add_item_numerado(doc, "Selecciona el cliente de la lista.")
add_item_numerado(doc, "Agrega uno o más equipos con su período de renta.")
add_item_numerado(doc, "El sistema calcula el total estimado basado en los precios de renta.")
add_item_numerado(doc, "El actor confirma la solicitud.")
add_item_numerado(doc, "El sistema crea la solicitud con estado Pendiente y solicitadoPor = usuario actual.")
add_item_numerado(doc, 'La solicitud aparece en el dashboard del administrador con badge "pendiente".')

add_seccion(doc, "Flujo Alterno")
add_item_lista(doc, "FA1 – Sin equipos seleccionados: El sistema requiere al menos un equipo.")
add_item_lista(doc, "FA2 – Equipo ya en renta: El sistema muestra advertencia de disponibilidad.")

add_seccion(doc, "Postcondiciones")
add_item_lista(doc, "La solicitud queda registrada con estado Pendiente.")
add_item_lista(doc, "El administrador y la secretaria pueden visualizarla para aprobar o rechazar.")

add_separador(doc)

# CU-19
add_titulo_cu(doc, "CU-19: Aprobar Solicitud de Renta")
add_campo(doc, "Nombre", "Aprobar Solicitud de Renta")
add_campo(doc, "Código", "CU-19")
add_campo(doc, "Actor principal", "Administrador, Secretaria")
add_campo(doc, "Descripción", "El actor revisa y aprueba una solicitud de renta pendiente, activando el contrato de renta.")

add_seccion(doc, "Precondiciones")
add_item_lista(doc, "El actor debe estar autenticado con rol admin o secretaria.")
add_item_lista(doc, "La solicitud debe existir con estado Pendiente.")

add_seccion(doc, "Flujo Básico")
add_item_numerado(doc, 'El actor accede a la sección Solicitudes y filtra por "Pendientes".')
add_item_numerado(doc, 'Selecciona una solicitud y hace clic en "Ver detalle".')
add_item_numerado(doc, "El sistema muestra RentaModal con: cliente, equipos, período, total estimado.")
add_item_numerado(doc, 'El actor revisa los datos y hace clic en "Aprobar".')
add_item_numerado(doc, "El sistema cambia el estado a Aprobada.")
add_item_numerado(doc, "Los equipos involucrados se marcan como en renta.")
add_item_numerado(doc, "La renta aparece en la sección Rentas Activas.")
add_item_numerado(doc, "El sistema muestra toast de éxito.")

add_seccion(doc, "Flujo Alterno")
add_item_lista(doc, "FA1 – Equipo ya no disponible: El sistema muestra advertencia.")

add_seccion(doc, "Postcondiciones")
add_item_lista(doc, "La solicitud queda con estado Aprobada.")
add_item_lista(doc, "El contrato aparece en Rentas Activas con el conteo de días.")

add_separador(doc)

# CU-20
add_titulo_cu(doc, "CU-20: Rechazar Solicitud de Renta")
add_campo(doc, "Nombre", "Rechazar Solicitud de Renta")
add_campo(doc, "Código", "CU-20")
add_campo(doc, "Actor principal", "Administrador, Secretaria")
add_campo(doc, "Descripción", "El actor rechaza una solicitud de renta pendiente, indicando el motivo, sin activar contrato alguno.")

add_seccion(doc, "Precondiciones")
add_item_lista(doc, "El actor debe estar autenticado con rol admin o secretaria.")
add_item_lista(doc, "La solicitud debe existir con estado Pendiente.")

add_seccion(doc, "Flujo Básico")
add_item_numerado(doc, "El actor localiza la solicitud pendiente.")
add_item_numerado(doc, 'Hace clic en "Rechazar".')
add_item_numerado(doc, "El sistema cambia el estado a Rechazada.")
add_item_numerado(doc, "Los equipos permanecen disponibles.")
add_item_numerado(doc, "El sistema muestra toast de confirmación.")

add_seccion(doc, "Flujo Alterno")
add_item_lista(doc, "FA1 – Solicitud ya procesada: El sistema muestra error de estado inválido.")

add_seccion(doc, "Postcondiciones")
add_item_lista(doc, "La solicitud queda con estado Rechazada.")
add_item_lista(doc, "Aparece en el historial de solicitudes rechazadas.")

add_separador(doc)

# CU-21
add_titulo_cu(doc, "CU-21: Ver Rentas Activas")
add_campo(doc, "Nombre", "Ver Rentas Activas")
add_campo(doc, "Código", "CU-21")
add_campo(doc, "Actor principal", "Administrador, Secretaria, Colaborador")
add_campo(doc, "Descripción", "El actor consulta la lista de contratos de renta activos con su progreso de días y datos del cliente.")

add_seccion(doc, "Precondiciones")
add_item_lista(doc, "El actor debe estar autenticado con cualquier rol permitido.")

add_seccion(doc, "Flujo Básico")
add_item_numerado(doc, "El actor accede a la sección Rentas Activas.")
add_item_numerado(doc, "El sistema muestra tarjetas de resumen: contratos activos, equipos en campo, ingresos proyectados.")
add_item_numerado(doc, "El sistema muestra la tabla de contratos con: ID, cliente, equipos, fecha inicio, fecha vencimiento, días restantes, total.")
add_item_numerado(doc, "Cada contrato incluye una barra de progreso visual (días usados vs restantes).")
add_item_numerado(doc, 'El actor puede hacer clic en "Ver detalle" para abrir RentaModal.')

add_seccion(doc, "Flujo Alterno")
add_item_lista(doc, "FA1 – Sin rentas activas: El sistema muestra mensaje de lista vacía.")

add_seccion(doc, "Postcondiciones")
add_item_lista(doc, "Ningún cambio de estado; es una operación de consulta.")

add_separador(doc)

# CU-22
add_titulo_cu(doc, "CU-22: Gestionar Renta Vencida")
add_campo(doc, "Nombre", "Gestionar Renta Vencida")
add_campo(doc, "Código", "CU-22")
add_campo(doc, "Actor principal", "Administrador")
add_campo(doc, "Descripción", "El administrador gestiona las rentas que han superado su fecha de vencimiento, pudiendo notificar al cliente o registrar la devolución del equipo.")

add_seccion(doc, "Precondiciones")
add_item_lista(doc, "El actor debe estar autenticado con rol admin.")
add_item_lista(doc, "Deben existir rentas con fecha de vencimiento superada.")

add_seccion(doc, "Flujo Básico")
add_item_numerado(doc, "El actor accede a la sección Rentas Vencidas.")
add_item_numerado(doc, "El sistema muestra el banner de alerta con el total de rentas vencidas.")
add_item_numerado(doc, "Para cada renta vencida se muestra: contrato, cliente con teléfono, equipos, fecha de vencimiento, días vencida y recargo estimado.")
add_item_numerado(doc, "El actor puede notificar individualmente a un cliente, notificar masivamente a todos (Notificar a todos), o marcar una renta como devuelta (Devuelto), lo que la mueve al historial con estado Dev. tardía.")

add_seccion(doc, "Flujo Alterno")
add_item_lista(doc, "FA1 – Sin rentas vencidas: La sección muestra lista vacía.")

add_seccion(doc, "Postcondiciones")
add_item_lista(doc, 'Si se marca como "Devuelto": la renta pasa al historial y los equipos quedan disponibles.')
add_item_lista(doc, "Los recargos calculados quedan visibles para cobro.")

add_separador(doc)

# CU-23
add_titulo_cu(doc, "CU-23: Consultar Historial de Rentas")
add_campo(doc, "Nombre", "Consultar Historial de Rentas")
add_campo(doc, "Código", "CU-23")
add_campo(doc, "Actor principal", "Administrador, Secretaria")
add_campo(doc, "Descripción", "El actor consulta el historial completo de contratos de renta con filtros y exportación a CSV.")

add_seccion(doc, "Precondiciones")
add_item_lista(doc, "El actor debe estar autenticado con rol admin o secretaria.")

add_seccion(doc, "Flujo Básico")
add_item_numerado(doc, "El actor accede a la sección Historial de rentas.")
add_item_numerado(doc, "El sistema muestra la tabla con: ID, cliente, equipos, período, total cobrado, estado, encargado.")
add_item_numerado(doc, "El actor puede filtrar por: texto de búsqueda, estado (Completada, Dev. tardía, Vencida, Rechazada), rango de fechas.")
add_item_numerado(doc, "El actor puede exportar los resultados a CSV con el botón de exportación.")

add_seccion(doc, "Flujo Alterno")
add_item_lista(doc, "FA1 – Sin resultados: El sistema muestra mensaje de lista vacía.")

add_seccion(doc, "Postcondiciones")
add_item_lista(doc, "Ningún cambio de estado; es una operación de consulta/exportación.")

add_separador(doc)

# CU-24
add_titulo_cu(doc, "CU-24: Consultar Bitácora de Cambios")
add_campo(doc, "Nombre", "Consultar Bitácora de Cambios")
add_campo(doc, "Código", "CU-24")
add_campo(doc, "Actor principal", "Administrador")
add_campo(doc, "Descripción", "El administrador consulta el registro de auditoría con todos los cambios realizados a equipos y usuarios, incluyendo valores anteriores y nuevos.")

add_seccion(doc, "Precondiciones")
add_item_lista(doc, "El actor debe estar autenticado con rol admin.")

add_seccion(doc, "Flujo Básico")
add_item_numerado(doc, "El actor accede a la sección Bitácoras.")
add_item_numerado(doc, "El sistema muestra estadísticas: total de registros, cambios hoy, cambios en equipos, cambios en usuarios.")
add_item_numerado(doc, "La tabla muestra: fecha/hora, módulo, entidad, campo modificado, valor anterior, valor nuevo, realizadoPor.")
add_item_numerado(doc, "El actor puede filtrar por: módulo (Todos / Equipos / Usuarios), búsqueda por nombre, campo, usuario o valores.")

add_seccion(doc, "Flujo Alterno")
add_item_lista(doc, "FA1 – Sin registros: El sistema muestra tabla vacía.")

add_seccion(doc, "Postcondiciones")
add_item_lista(doc, "Ningún cambio de estado; es una operación de solo lectura.")

add_separador(doc)

# CU-25
add_titulo_cu(doc, "CU-25: Ver Dashboard General")
add_campo(doc, "Nombre", "Ver Dashboard General")
add_campo(doc, "Código", "CU-25")
add_campo(doc, "Actor principal", "Administrador")
add_campo(doc, "Descripción", "El administrador visualiza un resumen ejecutivo del estado del sistema con KPIs, actividad reciente y próximas fechas de vencimiento.")

add_seccion(doc, "Precondiciones")
add_item_lista(doc, "El actor debe estar autenticado con rol admin.")

add_seccion(doc, "Flujo Básico")
add_item_numerado(doc, "El actor accede al sistema; la vista inicial es el Dashboard.")
add_item_numerado(doc, "El sistema muestra 4 KPI cards: Solicitudes pendientes, Rentas activas, Rentas vencidas, Ingresos del mes.")
add_item_numerado(doc, "El sistema muestra 5 mini-stats: Equipos en catálogo, disponibles, en renta, clientes registrados, usuarios del sistema.")
add_item_numerado(doc, "El sistema muestra el feed de actividad reciente: últimas 5 actividades con estado.")
add_item_numerado(doc, "El sistema muestra próximas a vencer: siguientes 4 contratos con indicador de urgencia por color.")

add_seccion(doc, "Flujo Alterno")
add_item_lista(doc, "FA1 – Sin datos: Las tarjetas muestran cero y los feeds aparecen vacíos.")

add_seccion(doc, "Postcondiciones")
add_item_lista(doc, "Ningún cambio de estado; es una operación de consulta en tiempo real.")

# =========================================================
# GUARDAR
# =========================================================
output_path = "/home/henryo/Escritorio/Proyecto_Tunsa/diagramas/Casos_de_Uso_Tunsa.docx"
doc.save(output_path)
print(f"Documento generado: {output_path}")
